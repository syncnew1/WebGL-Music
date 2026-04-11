from docx import Document
import re
from pathlib import Path

src = Path(r'c:\Users\zxq_1\Desktop\小米\WebGL Music\knowledge\毕业论文初稿.docx')
out = Path(r'c:\Users\zxq_1\Desktop\小米\WebGL Music\knowledge\毕业论文初稿-项目对齐优化版.docx')

doc = Document(str(src))

heading_re = re.compile(r'^(第[一二三四五六七八九十]+章|\d+(?:\.\d+)+)')

def para_text(p):
    return ''.join(run.text for run in p.runs)

def clear_para(p):
    for run in p.runs:
        run.text = ''

def find_heading_idx(title):
    for i, p in enumerate(doc.paragraphs):
        if para_text(p).strip() == title:
            return i
    raise ValueError(f'heading not found: {title}')

def next_heading_idx(start):
    for i in range(start + 1, len(doc.paragraphs)):
        t = para_text(doc.paragraphs[i]).strip()
        if heading_re.match(t):
            return i
    return len(doc.paragraphs)

def replace_section(title, new_paragraphs):
    start = find_heading_idx(title)
    end = next_heading_idx(start)
    first_body = start + 1
    if first_body >= len(doc.paragraphs):
        return
    clear_para(doc.paragraphs[first_body])
    doc.paragraphs[first_body].add_run(new_paragraphs[0])
    parent = doc.paragraphs[first_body]._element.getparent()
    insert_after = doc.paragraphs[first_body]._element
    for text in new_paragraphs[1:]:
        new_p = doc.add_paragraph(text)
        parent.insert(parent.index(insert_after) + 1, new_p._element)
        insert_after = new_p._element
    for idx in range(first_body + 1, end):
        clear_para(doc.paragraphs[idx])

replacements = {
    '2.3 WebGL 与 Three.js': [
        'WebGL 是浏览器端图形渲染标准，能够直接利用 GPU 完成高频图形绘制。相较于普通 DOM 动画和 Canvas 2D，WebGL 更适合处理频谱粒子、动态环形结构、脉冲扩散以及大规模像素级视觉反馈等需要持续刷新和较高并行计算能力的场景。Three.js 则是对 WebGL 的高级封装，可以在场景组织、着色器管理和渲染流程抽象方面降低开发复杂度。',
        '结合本项目的实际实现来看，系统的可视化方案采用了“Web Audio API + WebGL + Canvas 2D”混合渲染路径：一方面，封面脉冲模式与频谱环模式通过自定义 WebGL 渲染器完成高性能绘制；另一方面，频段能量条、和声轮盘、空间分布雷达图和乐器识别列表等分析组件则通过 Canvas 2D 与普通 React 组件进行组合呈现。该实现方式既保留了 WebGL 在动态视觉效果方面的性能优势，又降低了所有可视化内容统一三维化所带来的复杂度，更适合课程设计和浏览器端工程落地。'
    ],
    '2.4 Supabase': [
        'Supabase 是一种后端即服务平台，底层基于 PostgreSQL，向前端应用提供认证、数据库、对象存储以及访问控制能力。与传统自建后端相比，Supabase 能够显著降低中小型 Web 系统在用户管理、表结构维护、资源存储和权限配置方面的开发成本。',
        '在本系统中，Supabase 主要承担以下职责：一是完成用户注册、登录及个人资料初始化；二是存储歌曲、歌单、歌单-歌曲关联、播放历史、分享令牌和用户资料等结构化数据；三是通过对象存储保存音频文件和封面资源；四是利用 PostgreSQL 行级安全策略对用户私有数据实施访问隔离。项目代码中通过 songs、playlists、playlist_songs、playback_history、share_tokens 和 profiles 等数据表组织业务数据，并结合 signed URL、公共 URL 和对象下载回退机制实现音频资源访问。'
    ],
    '3.1.5 音频可视化功能': [
        '（一）封面脉冲可视化：根据低频能量、节拍强度和时间参数驱动中心区域的脉冲扩散与光晕变化；',
        '（二）频谱环展示：将频谱幅值映射为环形点阵与半径变化，用于表现音乐强弱起伏；',
        '（三）频谱分析视图：以连续频谱柱状/波形方式反映当前频域分布；',
        '（四）频段能量分析：对超低频、低频、中频、高频等频带进行划分并动态显示能量变化；',
        '（五）和声分析与空间分布：展示当前主要音高、和声结构和空间化分布状态；',
        '（六）乐器识别与动态显示：基于频段特征估算活跃乐器，并以列表和能量条形式呈现其活跃程度。'
    ],
    '4.1 系统总体架构设计': [
        '本系统采用前后端分离架构，整体由表示层、状态与业务逻辑层、音频分析与渲染层以及云端数据层四部分组成。该架构既保证了前端界面的交互灵活性，也使播放控制、歌词同步、可视化分析和数据访问等职责能够相对独立地组织。系统总体架构如图4-1所示。',
        '在表示层中，系统基于 React 构建首页、搜索页、音乐库、歌单详情页、个人中心、登录注册页以及歌词抽屉和播放器底栏等界面组件。状态与业务逻辑层通过多个 Provider 对系统核心状态进行分层管理，其中 PlayerProvider 负责音频播放链路与队列控制，DataProvider 负责歌曲、歌单、歌词与历史记录等数据操作，LayoutProvider 负责界面布局展开收起，VisualizerProvider 负责可视化模式、主题和参数切换。',
        '音频分析与渲染层建立在 Web Audio API 与图形渲染模块之上：PlayerProvider 中构建 AudioContext、GainNode、DynamicsCompressorNode、BiquadFilterNode 与 AnalyserNode 组成的音频节点图；AudioAnalyzer 对实时采样结果进行频段、响度、和声、瞬态和乐器特征分析；随后由 WebGL 与 Canvas 2D 组件完成可视化呈现。数据层则依托 Supabase 提供认证、数据库与对象存储能力，并结合 PostgreSQL 行级安全策略实现多用户数据隔离与访问控制。'
    ],
    '4.2.1 播放器模块': [
        '播放器模块是系统核心业务模块，主要由 PlayerProvider 统一管理。该模块负责音频元素生命周期、播放状态维护、音量控制、播放模式切换、播放队列管理、进度更新和异常处理，并对外向歌词模块和可视化模块提供统一的播放时间与分析节点。系统支持播放、暂停、上一首、下一首、拖动进度、静音、音量渐变以及顺序播放、单曲循环、随机播放等功能。'
    ],
    '4.2.2 数据管理模块': [
        '数据管理模块主要由 DataProvider 负责，封装了本地缓存与 Supabase 云端数据访问逻辑。模块内部统一处理歌曲上传、歌曲删除、歌单创建、歌单歌曲关联维护、搜索、历史记录写入、歌词更新、封面和歌词自动补全以及“喜欢”歌单维护等操作。对于无云端配置场景，系统可退化使用 localStorage 保存基础数据；对于云端模式，系统通过 Supabase 数据表和对象存储实现跨会话数据持久化。'
    ],
    '4.2.3 可视化分析模块': [
        '可视化分析模块由 VisualizerProvider、AudioAnalyzer、InsightDashboard、WebGLModeCanvas 及多种可视化组件共同构成。模块先通过 AnalyserNode 获取音频时域与频域数据，再由 AudioAnalyzer 计算频段能量、响度、谱质心、谱流、节拍强度、和声信息和乐器活跃度，最后根据模式切换结果分别驱动封面脉冲、频谱环、频谱视图、空间分布雷达图、和声轮盘和乐器列表等界面。为避免瞬时抖动，模块内部引入了平滑算法、短时保留机制与动画缓动策略。'
    ],
    '4.2.4 布局管理模块': [
        '布局管理模块主要由 LayoutProvider 与播放器侧边面板控制逻辑组成，用于协调左侧导航栏、中间内容区、右侧歌词/队列面板以及中部可视化面板的展开与收起。该模块保证在不同视图模式下界面布局能够保持一致，并使播放器、歌词抽屉和分析仪表盘之间形成联动。'
    ],
    '4.3 数据库设计': [
        '根据项目实际实现，系统数据库的核心数据表包括 songs、playlists、playlist_songs、playback_history、share_tokens 和 profiles。songs 表用于存储歌曲标题、歌手、专辑、标签、音频路径、封面路径、歌词文本及所有者信息；playlists 表用于存储歌单名称、描述、公开状态及所有者；playlist_songs 表维护歌单与歌曲之间的多对多关系；playback_history 表记录用户播放历史；share_tokens 表用于私密分享歌单链接；profiles 表用于保存用户基本资料。',
        '数据库之外，系统还依赖 Supabase 对象存储中的 audio 和 covers 两个 bucket，用于保存音频资源和封面文件。为保障安全性，系统在 songs、playlists、playlist_songs、playback_history、share_tokens 和 profiles 等表上启用了行级安全策略，并通过 owner_id 或 user_id 与 auth.uid() 的匹配关系限制访问范围。对于公开歌曲与公开歌单，系统还额外提供选择性读取策略，从而兼顾资源共享与数据隔离。'
    ],
    '4.4.2 可视化渲染流程': [
        '在可视化渲染流程中，系统首先由 AnalyserNode 周期性采集音频频域与时域数据。随后，AudioAnalyzer 对数据进行进一步处理，计算各频带能量、RMS、LUFS 近似值、谱质心、谱流、平滑后的低中高频能量、节拍强度以及和声和乐器特征。最后，根据当前可视化模式选择不同渲染路径：封面脉冲模式使用 WebGL 着色渲染脉冲扩散与光晕，频谱环模式使用 WebGL 点阵绘制环形频谱，频谱模式则以 Canvas 方式展示频谱波形，其余分析图表使用 Canvas 2D 与普通组件呈现。该流程实现了分析逻辑与绘制逻辑的解耦，有利于后续扩展更多可视化模式。'
    ],
    '5.1 开发环境与项目结构': [
        '本系统前端采用 React 18、TypeScript 与 Vite 构建单页应用，样式层结合 Tailwind CSS 与自定义主题变量实现统一视觉风格。音频分析依托 Web Audio API，图形渲染采用 WebGL 与 Canvas 2D 组合方式，后端数据服务由 Supabase 提供。',
        '从项目结构来看，src/components 目录下主要存放播放器、歌词、导航栏、内容卡片及可视化仪表盘等界面组件；src/pages 用于组织首页、搜索、音乐库、歌单、登录注册和工具页等页面；src/providers 中集中管理 PlayerProvider、DataProvider、LayoutProvider 与 VisualizerProvider 四类全局状态；src/visualizer 中实现 AudioAnalyzer 以及 gl 渲染器；supabase 目录保存数据库表结构与权限策略脚本。该目录组织方式有利于界面组件、状态逻辑、音频分析与数据访问分层维护。'
    ],
    '5.2 音频播放模块实现': [
        '播放器模块通过 PlayerProvider 统一管理全局播放状态，是整个系统的核心业务模块之一。系统初始化时创建 HTMLAudioElement，用户触发播放操作后再按需构建音频处理链路，以减少不必要的资源占用并提高浏览器兼容性。',
        '在具体实现中，音频链路主要包括 MediaElementAudioSourceNode、GainNode、DynamicsCompressorNode、BiquadFilterNode 和 AnalyserNode。GainNode 负责音量控制与渐变调节；DynamicsCompressorNode 用于限幅与动态范围压缩；BiquadFilterNode 通过高通滤波削弱部分低频干扰；AnalyserNode 则为频谱分析和可视化提供实时数据源。播放过程中，系统监听 timeupdate、durationchange、ended 和 error 等事件维护当前时间、总时长、播放状态与错误提示，并根据顺序播放、单曲循环和随机播放模式自动切换曲目。',
        '此外，为提升云端资源访问的稳定性，播放器在处理 songs 表中的 storage_path 与 url 字段时，支持 signed URL、公共 URL 和对象下载回退三种资源解析方式，并对签名地址进行短期缓存，以降低重复请求带来的性能开销。'
    ],
    '5.3 歌词模块实现': [
        '歌词模块由 LyricsPanel 与 DataProvider 中的歌词维护逻辑共同实现。系统支持普通文本歌词与 LRC 时间标签歌词两种形式：在解析过程中，程序首先按行拆分歌词文本，再提取每行中的时间标签和正文内容，最终构建按时间排序的歌词数组。播放过程中，系统依据 currentTime 实时计算当前活动歌词行，并通过自动滚动和高亮策略将当前行定位到可视区域中心。',
        '从交互能力上看，当前实现不仅支持同步显示，还支持歌词编辑、时间标签插入、歌词保存以及点击歌词跳转播放位置等操作。为适应不同来源的歌词文本，项目还引入了繁体转简体处理逻辑，用于统一歌词显示形式。相较于仅实现单纯滚动显示的方案，本系统的歌词模块更强调可编辑性、同步性和用户交互体验。'
    ],
    '5.4 可视化分析模块实现': [
        '可视化分析模块是本系统区别于传统音乐播放器的核心特色之一。系统通过 AnalyserNode 获取原始采样数据后，交由 AudioAnalyzer 进一步计算频段能量、平滑后的低中高频强度、RMS、LUFS 近似值、谱质心、谱流、节拍强度、和声集合和乐器活跃度，并将这些分析结果传递给可视化界面。',
        '在具体呈现形式上，系统实现了三类主模式与多种辅助分析视图。主模式包括：封面脉冲模式，用于通过低频与节拍驱动中心区域脉冲扩散；频谱环模式，用于以环形点阵表现频域强度分布；频谱模式，用于直接显示连续频谱变化。辅助分析视图包括频段能量条、和声轮盘、空间分布雷达图和乐器识别列表，用于展示不同维度的分析结果。',
        '为增强视觉稳定性，项目在分析和渲染过程中引入了指数平滑、短时保留和渐变过渡机制。例如，AudioAnalyzer 对乐器能量采用“快速攻击、缓慢衰减”的平滑策略，并通过 hold 帧数控制避免乐器标签频繁闪烁；频段条和可视化组件则通过插值和过渡动画减少瞬时抖动。这些策略有效改善了浏览器端高频渲染场景下的观感。'
    ],
    '5.5 后端数据模块实现': [
        '系统后端数据能力由 Supabase 提供。用户登录后，前端通过 SDK 拉取 songs、playlists 和 playlist_songs 等表中的数据，并在本地状态中重建歌曲列表和歌单关系；同时，播放历史通过 playback_history 表进行记录，用户资料通过 profiles 表进行维护。对于歌单分享功能，系统在 share_tokens 表中存储令牌哈希和过期时间，并通过数据库函数提供基于令牌的共享歌单访问能力。',
        '在资源管理方面，音频文件与封面文件分别存储于 Supabase 的 audio 与 covers bucket 中。上传歌曲时，系统会先尝试读取本地音频标签信息，如标题、歌手、专辑和封面，再将音频文件上传到对象存储并写入 songs 表。播放时优先根据 storage_path 生成 signed URL，必要时再退回公共 URL 或对象下载方式，以保证资源访问可靠性。对于未启用云端环境的场景，DataProvider 还保留了 localStorage 兜底逻辑，从而提升系统的可运行性与适应性。'
    ],
    '6.3 性能测试': [
        '针对音频分析与可视化部分，系统重点从页面响应速度、播放过程流畅性、可视化刷新稳定性以及音画同步表现等方面开展性能测试。结合项目实现特点，测试过程中主要关注以下内容：一是初次进入页面时组件加载与音频节点初始化是否及时；二是在持续播放状态下，歌词高亮、播放进度与可视化界面是否能够保持同步；三是在多种可视化模式切换时，浏览器界面是否出现明显卡顿、抖动或资源泄漏；四是在加载云端歌曲资源时，signed URL 缓存策略是否能够减少重复请求。',
        '测试结果表明，在将高频变化逻辑集中到 AudioAnalyzer、WebGL 渲染器和局部可视化组件后，系统能够避免大量业务组件参与高频重渲染，从而提高整体流畅度。与此同时，播放器中对音频节点的按需初始化、数据模块中的本地缓存与签名地址缓存，以及可视化模块中的频谱平滑和过渡动画，均对性能稳定性产生了积极作用。总体来看，在常规桌面浏览器环境下，系统能够较稳定地完成音频播放、歌词同步和多模式可视化展示。'
    ],
    '6.4 问题分析': [
        '在系统开发与测试过程中，主要暴露出以下几个方面的问题：',
        '1、音频播放与可视化渲染之间存在轻微同步偏差，尤其在歌曲切换、浏览器标签页切换或资源延迟加载时更为明显；',
        '2、乐器识别和频段分析结果在阈值附近波动较大，若不进行平滑处理，容易引发界面元素闪动；',
        '3、歌词来源存在时间标签格式不统一、繁简体混杂和部分歌词缺失等问题，影响显示一致性；',
        '4、云端资源访问依赖对象存储权限与签名地址，若权限策略或资源路径配置不当，可能导致部分歌曲无法直接播放。',
        '针对上述问题，系统分别采用了时间偏移修正、频谱与乐器能量平滑、短时保留机制、繁体转简体、歌词编辑保存以及 signed URL 缓存与回退访问策略等方法进行改进。整体来看，这些优化措施有效提升了系统的同步性、稳定性和容错性，但在更高精度的音乐信息检索、移动端渲染优化和大规模数据场景下仍有进一步完善空间。'
    ],
    '7.2 展望': [
        '尽管系统已基本完成设计目标，但仍有进一步优化空间：',
        '1. 引入更高精度的音乐信息检索或机器学习方法，提高乐器识别、节拍检测与和声分析准确率；',
        '2. 在现有封面脉冲、频谱环和频谱模式基础上扩展更多视觉效果，并完善主题系统与参数配置能力；',
        '3. 优化移动端适配与多终端交互体验，提升不同屏幕尺寸下的布局与渲染表现；',
        '4. 在云端数据能力基础上扩展推荐、收藏、社交分享和个性化画像等功能；',
        '5. 继续改进歌词模块，在保证稳定性的前提下研究更细粒度的逐字同步与音画联动机制。',
        '未来，可在本系统基础上进一步提升其智能分析能力、可视化表现力和实际应用场景适配能力，使其在在线音乐平台、数字媒体展示和交互艺术创作等方向发挥更大价值。'
    ]
}

for title, paras in replacements.items():
    replace_section(title, paras)

doc.save(str(out))
print(str(out))
