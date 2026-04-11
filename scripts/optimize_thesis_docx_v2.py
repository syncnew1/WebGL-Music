from docx import Document
import re
from pathlib import Path

src = Path(r'c:\Users\zxq_1\Desktop\小米\WebGL Music\knowledge\毕业论文初稿-项目对齐优化版.docx')
out = Path(r'c:\Users\zxq_1\Desktop\小米\WebGL Music\knowledge\毕业论文初稿-项目对齐优化版-v2.docx')

doc = Document(str(src))
heading_re = re.compile(r'^(第[一二三四五六七八九十]+章|\d+(?:\.\d+)+)')

def para_text(p):
    return ''.join(run.text for run in p.runs)

def clear_para(p):
    for run in p.runs:
        run.text = ''

def set_para(p, text):
    clear_para(p)
    p.add_run(text)

def find_heading_idx(title):
    for i, p in enumerate(doc.paragraphs):
        if para_text(p).strip() == title:
            return i
    raise ValueError(title)

def next_heading_idx(start):
    for i in range(start + 1, len(doc.paragraphs)):
        t = para_text(doc.paragraphs[i]).strip()
        if heading_re.match(t):
            return i
    return len(doc.paragraphs)

def replace_section(title, new_paragraphs):
    start = find_heading_idx(title)
    end = next_heading_idx(start)
    body_start = start + 1
    set_para(doc.paragraphs[body_start], new_paragraphs[0])
    parent = doc.paragraphs[body_start]._element.getparent()
    insert_after = doc.paragraphs[body_start]._element
    for text in new_paragraphs[1:]:
      new_p = doc.add_paragraph(text)
      parent.insert(parent.index(insert_after) + 1, new_p._element)
      insert_after = new_p._element
    for idx in range(body_start + 1, end):
      clear_para(doc.paragraphs[idx])

def replace_exact(old, new):
    for p in doc.paragraphs:
        if para_text(p).strip() == old:
            set_para(p, new)
            return True
    return False

# 图示占位改为可直接画图的真实说明
image_map = {
    '【图片占位符：系统用例图】': '制图说明：系统用例图建议包含“游客/注册用户”两个参与者。游客可执行注册、登录、浏览公开歌曲等操作；注册用户在此基础上还可执行播放歌曲、暂停/切歌、搜索歌曲、创建歌单、添加或移除歌曲、查看歌词、编辑歌词、查看可视化面板、查看播放历史、批量上传歌曲和分享歌单等用例。图中以播放器系统为边界框，参与者置于左侧，用椭圆表示各业务用例，并用关联线连接对应角色。',
    '【图片占位符：系统总体架构图】': '制图说明：系统总体架构图建议采用四层结构。最上层为表示层，包括 TopBar、NavSidebar、PlayerControls、LyricsPanel、InsightDashboard 及各页面；第二层为状态与业务逻辑层，包括 PlayerProvider、DataProvider、LayoutProvider、VisualizerProvider；第三层为音频分析与渲染层，包括 AudioContext、GainNode、DynamicsCompressorNode、BiquadFilterNode、AnalyserNode、AudioAnalyzer、WebGL 渲染器和 Canvas 分析组件；最下层为数据层，包括 Supabase Auth、PostgreSQL、Storage(audio/covers)。各层之间用箭头标明调用与数据流向。',
    '【图片占位符：系统功能模块图】': '制图说明：系统功能模块图建议以“音乐可视化系统”为中心，向外分为播放器模块、数据管理模块、歌词模块、可视化分析模块、布局管理模块和用户管理模块。播放器模块下细分播放控制、进度控制、音量控制、播放模式；数据模块下细分歌曲管理、歌单管理、历史记录、资源加载；可视化分析模块下细分频段分析、和声分析、乐器识别、空间分布、模式切换；用户模块下细分登录注册、个人资料与权限控制。',
    '【图片占位符：数据库 E-R 图】': '制图说明：数据库 E-R 图应至少包含 profiles、songs、playlists、playlist_songs、playback_history、share_tokens 六个实体。profiles 与 songs/playlists/playback_history 通过用户标识关联；playlists 与 songs 之间通过 playlist_songs 构成多对多关系；share_tokens 与 playlists 通过 resource_id 关联。图中应标注主键、外键和主要字段，并体现 songs 与 playlists 的公开/私有属性。',
    '【图片占位符：系统业务流程图】': '制图说明：系统业务流程图建议按纵向流程绘制，主线为“用户选择歌曲并点击播放”→“解析歌曲地址(url/storage_path)”→“初始化 AudioContext 与音频节点”→“开始播放音频”→“AnalyserNode 采集时域/频域数据”。随后分为两条支路：一条进入“歌词时间定位与滚动更新”，另一条进入“AudioAnalyzer 计算频段/和声/乐器/节拍”→“WebGL/Canvas 组件渲染”。最后汇聚到“界面实时更新（播放进度、歌词、可视化）”。',
    '【图片占位符：页面布局结构图】': '制图说明：页面布局结构图建议使用框架示意图展示界面区域关系。顶部为 TopBar，左侧为 NavSidebar，中部为内容区（Home/Search/Library/Playlist 等页面），右侧为可切换侧边面板（queue / lyrics），底部为 PlayerControls，中间可叠加可视化仪表盘区域。图中建议使用箭头说明左侧导航切换页面、底部播放器驱动右侧歌词与中部可视化联动的关系。',
    '【图片占位符：音频播放模块流程图】': '制图说明：音频播放模块流程图可采用泳道式或纵向流程图。主要步骤为“用户点击播放”→“读取当前歌曲信息”→“解析 storage_path 或 url”→“创建/恢复 AudioContext”→“连接 GainNode、Compressor、High-pass Filter、AnalyserNode”→“调用 audio.play()”→“监听 timeupdate/durationchange/ended/error 事件”→“更新进度、时长、模式切歌与错误提示”。图中可补充 signed URL 缓存与回退获取的分支。',
    '【图片占位符：歌词同步模块流程图】': '制图说明：歌词同步模块流程图建议包含“读取歌词文本”→“解析 LRC 时间标签”→“生成按时间排序的歌词序列”→“根据 currentTime 定位当前行”→“执行高亮和自动滚动更新”→“用户点击歌词可触发 seek 跳转”。如需体现编辑功能，可在旁路增加“插入当前时间标签”和“保存歌词到云端”两个步骤。',
    '【图片占位符：音频分析处理流程图】': '制图说明：音频分析处理流程图建议绘制为数据处理链。输入为 AnalyserNode 频域/时域数据；中间依次为“计算频段能量”“计算 RMS/LUFS/谱质心/谱流”“平滑低中高频”“节拍检测”“和声检测”“乐器识别”；输出为“频段能量条”“和声轮盘”“空间分布雷达图”“乐器识别列表”“封面脉冲/频谱环/频谱视图”。',
    '【图片占位符：可视化分析界面截图】': '截图说明：建议截取系统的可视化分析主界面，画面中应同时体现顶部模式控制区、主预览区域（封面脉冲或频谱环）、统计信息区以及下方的声场分布、和声分析、频段能量和乐器识别等组件，确保截图能够反映系统“分析 + 可视化”一体化特征。',
    '【图片占位符：系统主要功能界面展示图】': '截图说明：建议使用四宫格拼图展示系统主要界面，其中(a)为主页或播放主界面，(b)为歌词显示与编辑界面，(c)为音乐库或歌单管理界面，(d)为可视化分析界面。每张子图下方应标注对应功能名称，用于说明系统主要业务模块均已实现。'
}
for old, new in image_map.items():
    replace_exact(old, new)

# 继续增强正文
replace_section('3.2 非功能需求分析', [
    '除功能完整性外，系统还需要在性能、安全、可维护性和可用性等方面满足 Web 多媒体应用的基本要求。由于本课题同时涉及音频播放、实时分析和动态图形渲染，因此非功能需求对系统最终质量具有重要影响。',
    '3.2.1 性能需求',
    '系统应在普通桌面浏览器环境下保持较为流畅的交互体验，音频播放、歌词滚动与可视化渲染之间应具有较低延迟。在连续播放场景中，系统需尽量避免界面卡顿、画面撕裂或因高频状态更新导致的明显掉帧现象。对于可视化模块，应通过频谱平滑、渲染循环优化、节点按需初始化和组件状态隔离等手段维持较稳定的刷新效果。',
    '3.2.2 安全需求',
    '系统涉及用户账户、歌单、播放历史以及音频资源访问等数据内容，因此必须具备基本的数据安全控制能力。未授权用户不得访问其他用户的私有歌曲、私有歌单和播放历史；音频对象存储访问应受存储策略约束；分享链接应通过令牌和过期时间控制访问有效期。系统通过 Supabase Auth 与 PostgreSQL 行级安全策略实现上述安全要求。',
    '3.2.3 可维护性需求',
    '系统应具备清晰的模块划分、明确的状态边界以及较好的代码可读性，便于后续进行功能扩展、性能优化和问题定位。项目采用 Provider 分层管理方式，将播放逻辑、数据逻辑、布局逻辑和可视化控制逻辑分别封装，降低模块间耦合，并通过 TypeScript 类型约束提高维护效率。',
    '3.2.4 可用性需求',
    '界面布局应简洁直观，交互路径应符合用户对音乐播放器的常规认知。系统除保证播放、切歌、拖动进度、查看歌词等基础操作易用外，还应通过歌词高亮、播放状态反馈、动态可视化和统一主题风格增强使用体验。对于无云端配置场景，系统仍应能够以本地模式完成基础功能，以提高可用性和环境适应能力。'
])

replace_section('6.1 测试目标', [
    '系统测试主要从功能正确性、性能表现和用户体验三个方面进行，目标是验证系统是否满足毕业设计预期需求，并能够在浏览器环境中稳定完成音乐播放、歌词同步、数据交互和实时可视化。',
    '从功能层面，测试需要确认用户认证、音乐播放控制、歌单管理、歌词显示与编辑、搜索、历史记录以及可视化模式切换等核心业务是否可正常执行；从性能层面，需要关注页面加载、资源解析、音频节点初始化、歌词滚动和可视化刷新是否流畅；从交互体验层面，则主要考察界面反馈是否清晰、模块切换是否自然以及音画联动是否具有较好的沉浸感。'
])

replace_section('6.2 功能测试', [
    '功能测试主要围绕系统核心业务流程展开，包括用户认证、歌曲播放、歌词同步、数据管理和可视化展示等方面。测试过程中采用人工交互验证与界面观察相结合的方式，对每项核心功能进行逐项检查。',
    '1、用户登录与退出是否正常，未登录状态下是否限制上传等需要鉴权的操作；',
    '2、音乐播放控制是否完整可用，包括播放、暂停、上一首、下一首、拖动进度、静音与音量调节；',
    '3、歌词加载、同步滚动、点击跳转与保存编辑功能是否准确；',
    '4、歌单、历史记录、搜索以及歌曲上传删除功能是否正确；',
    '5、封面脉冲、频谱环、频谱等可视化模式是否能够随播放实时变化。',
    '由测试结果可以看出，系统的核心功能模块能够基本按预期运行，业务链路完整，用户能够在统一界面下完成从歌曲浏览、播放控制到歌词查看和可视化分析的完整操作流程。'
])

replace_section('6.4 问题分析', [
    '在系统开发与测试过程中，主要暴露出以下几个方面的问题：',
    '1、音频播放与可视化渲染之间存在轻微同步偏差，尤其在歌曲切换、浏览器标签页调度变化或资源加载延迟时更为明显；',
    '2、乐器识别和频段分析结果在阈值附近波动较大，如果直接驱动界面元素，容易出现标签闪动和能量条抖动；',
    '3、歌词来源的格式差异较大，不同歌曲可能存在时间标签精度不一致、歌词缺失、繁简体混杂等问题，影响同步显示效果；',
    '4、云端资源访问依赖对象存储权限和 signed URL 解析策略，若资源路径、权限策略或网络状态出现异常，可能造成部分歌曲无法即时播放。',
    '针对上述问题，系统分别采用了时间偏移修正、频谱平滑、乐器短时保留、状态拆分、歌词编辑补偿、繁体转简体以及签名地址缓存与回退访问等优化措施。测试结果表明，这些改进在一定程度上提升了系统的稳定性、同步效果和交互连续性，但在更高精度的音乐信息识别和更复杂终端环境中仍有进一步提升空间。'
])

doc.save(str(out))
print(str(out))
